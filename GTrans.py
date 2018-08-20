import docx
from googletrans import Translator
import time
import os

```
日语docx文档翻译脚本
版本：v1.0
```

doc_name = 'SAG.docx'
ts = Translator(service_urls=['translate.google.cn'])
doc = docx.Document(doc_name)
paras = doc.paragraphs
tables = doc.tables


def transPara(text):
    if ts.detect(text).lang == 'ja':
        if '表 ' in text:
            text = text.replace('表 ', '')
            text = '表 ' + ts.translate(text, dest='zh-cn').text
        else:
            text = ts.translate(text, dest='zh-cn').text
        print(text)
    return text

def transCell(text):
    if '#' in text and len(text) == 1:
        pass
    elif ts.detect(text).lang == 'ja':
        text = ts.translate(text, dest='zh-cn').text
        print(text)
    return text

# translate the paragraph

if os.path.exists('pn.txt'):
    with open('pn.txt', 'r') as f:
        pn = int(f.readline())
else:
    pn = 0

_pn = pn
for para in paras[pn:]:
    if len(para.text) > 0:
        try:
            para.text = transPara(para.text)
        except Exception as e:
            try:
                print('+' * 20 + ' New Translator ' + '+' * 20)
                ts = Translator(service_urls=['translate.google.cn'])
                para.text = transPara(para.text)
            except Exception as e:
                with open('pn.txt', 'w') as f:
                    f.write(str(_pn))
                doc.save('TS_' + doc_name)
                exit(1)
    _pn += 1
doc.save('TS_' + doc_name)



if os.path.exists('tn.txt'):
    with open('tn.txt', 'r') as f:
        tn = int(f.readline())
else:
    tn = 0
_tn = tn
for table in tables[tn:]:
    for row in table.rows:
        for cell in row.cells:
            if len(cell.text) > 0:
                try:
                    cell.text = transCell(cell.text)
                except Exception as e:
                    try:
                        ts = Translator(service_urls=['translate.google.cn'])
                        print('+' * 20 + ' New Translator ' + '+' * 20)
                        cell.text = transCell(cell.text)
                    except Exception as e:
                        with open('tn.txt', 'w') as f:
                            f.write(str(_tn))
                        doc.save('TS_' + doc_name)
                        exit(2)

    _tn += 1
doc.save('TS_' + doc_name)
